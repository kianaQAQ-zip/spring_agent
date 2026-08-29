"""多格式导出器（txt / md / docx / xlsx / pdf）。

统一入口 `export(fmt, text, blocks)` → `(bytes, content_type, file_ext)`。
- txt/md：纯文本直写（UTF-8）
- docx：python-docx 重建（标题层级 + 段落 + 表格）
- xlsx：openpyxl（正文按行入单元格；有 table 块时按块拆工作表）
- pdf ：fpdf2 + Windows 系统中文字体（无中文字体时返回友好错误，不产乱码）

损坏/不支持格式抛 `UnsupportedFormatError`（友好中文提示）。
"""
from __future__ import annotations

import io
import os
import re

from .models import ParseBlock
from .parsers import UnsupportedFormatError

SUPPORTED_FORMATS = ("txt", "md", "docx", "xlsx", "pdf")


def export(fmt: str, text: str, blocks: list[ParseBlock] | None = None) -> tuple[bytes, str, str]:
    f = (fmt or "").strip().lower()
    if not text and not blocks:
        raise UnsupportedFormatError("没有可导出的内容")
    if f == "txt":
        return (text or "").encode("utf-8"), "text/plain; charset=utf-8", "txt"
    if f == "md":
        return (text or "").encode("utf-8"), "text/markdown; charset=utf-8", "md"
    if f == "docx":
        return _export_docx(text, blocks), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    if f == "xlsx":
        return _export_xlsx(text, blocks), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"
    if f == "pdf":
        return _export_pdf(text, blocks), "application/pdf", "pdf"
    raise UnsupportedFormatError(f"不支持的导出格式：{fmt or '空'}（支持 {', '.join(SUPPORTED_FORMATS)}）")


def _heading_level(text: str) -> int:
    m = re.match(r"^(#{1,6})\s", text)
    return len(m.group(1)) if m else 0


def _iter_source(blocks: list[ParseBlock] | None, text: str):
    """统一迭代源：优先结构化 blocks，否则纯文本按空行分段落。"""
    if blocks:
        yield from blocks
    else:
        order = 0
        for para in re.split(r"\n\s*\n", text or ""):
            p = para.strip()
            if p:
                yield ParseBlock(block_type="text", text=p, bbox=None, page=1,
                                 reading_order=order, token_count=0)
                order += 1


# ---------------------------------------------------------------------------
# docx
# ---------------------------------------------------------------------------
def _export_docx(text: str, blocks: list[ParseBlock] | None) -> bytes:
    try:
        import docx
    except ImportError:
        raise UnsupportedFormatError("未安装 python-docx，无法导出 Word 文档") from None

    doc = docx.Document()
    for b in _iter_source(blocks, text):
        bt = b.block_type or "text"
        if bt == "heading":
            doc.add_heading(b.text.lstrip("#").strip(), level=max(1, min(9, _heading_level(b.text) or 1)))
        elif bt == "table":
            rows = [r.split(" | ") for r in b.text.split("\n") if r.strip()]
            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                for i, row in enumerate(rows):
                    for j in range(ncols):
                        table.cell(i, j).text = row[j] if j < len(row) else ""
                doc.add_paragraph()
        else:
            doc.add_paragraph(b.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# xlsx
# ---------------------------------------------------------------------------
def _export_xlsx(text: str, blocks: list[ParseBlock] | None) -> bytes:
    try:
        from openpyxl import Workbook
    except ImportError:
        raise UnsupportedFormatError("未安装 openpyxl，无法导出 Excel 表格") from None

    wb = Workbook()
    ws = wb.active
    ws.title = "导出内容"

    if blocks and any(b.block_type == "table" for b in blocks):
        # 结构化导出：每个 table 块一个工作表，非表格块放第一个 sheet
        first = True
        for b in blocks:
            if b.block_type == "table":
                if not first:
                    ws = wb.create_sheet()
                ws.title = f"Sheet{len(wb.worksheets)}"
                for row in b.text.split("\n"):
                    ws.append([c.strip() for c in row.split(" | ")])
                first = False
        return _save_wb(wb)
    else:
        for line in (text or "").split("\n"):
            if line.strip():
                ws.append([line])
        return _save_wb(wb)


def _save_wb(wb) -> bytes:
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# pdf
# ---------------------------------------------------------------------------
_CHINESE_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\simhei.ttf",   # 黑体
    r"C:\Windows\Fonts\msyh.ttf",     # 微软雅黑 Regular
    r"C:\Windows\Fonts\simsun.ttc",   # 宋体
    r"C:\Windows\Fonts\msyh.ttc",     # 微软雅黑
    r"C:\Windows\Fonts\simfang.ttf",  # 仿宋
]


def _find_chinese_font() -> str | None:
    for p in _CHINESE_FONT_CANDIDATES:
        if os.path.exists(p):
            return p
    # macOS / Linux 常见路径
    for p in ("/System/Library/Fonts/PingFang.ttc", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"):
        if os.path.exists(p):
            return p
    return None


def _export_pdf(text: str, blocks: list[ParseBlock] | None) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError:
        raise UnsupportedFormatError("未安装 fpdf2，无法导出 PDF") from None

    font_path = _find_chinese_font()
    if font_path is None:
        raise UnsupportedFormatError("导出 PDF 需要系统中文字体（未找到 simhei/msyh 等），已中止以避免乱码")

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("cjk", "", font_path)
    pdf.set_font("cjk", size=11)

    for b in _iter_source(blocks, text):
        bt = b.block_type or "text"
        content = b.text
        if bt == "heading":
            pdf.set_font("cjk", size=13)
            pdf.multi_cell(0, 7, content.lstrip("#").strip())
            pdf.set_font("cjk", size=11)
        elif bt == "table":
            for line in content.split("\n"):
                pdf.multi_cell(0, 6, line)
        else:
            pdf.multi_cell(0, 6, content)
        pdf.ln(1)

    return bytes(pdf.output())
