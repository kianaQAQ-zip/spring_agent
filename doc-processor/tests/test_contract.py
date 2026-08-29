"""doc-processor 接口契约测试（与 Java 侧 EvalRunner 解耦，可独立跑）。

验证 §10.2 契约：统一信封 {ok,data,error}、clean/parse 真实跑通、rerank 模型不可用时降级。
"""
from __future__ import annotations

import pytest
from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)

NOISY_TEXT = (
    "\ufeff【售后政策 V3】\n"
    "第一章 七天无理由退货.....12\n"
    "客户购买商品后，自签收之日起七日内，\n"
    "在商品完好不影响二次销售的前提下，可申请无理由退货。\n"
    "第一章 七天无理由退货.....12\n"
    "客服电话：13812345678，邮箱：service@example.com\n"
)


def test_health():
    r = client.get("/health")
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True
    assert body["data"]["status"] == "UP"


def test_clean_basic():
    r = client.post("/api/v1/clean", json={"text": NOISY_TEXT})
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True
    assert set(body.keys()) >= {"ok", "data", "error"}
    data = body["data"]
    assert "七天无理由退货" in data["cleaned_text"]
    assert 0.0 <= data["clean_score"] <= 1.0
    assert isinstance(data["removed_flags"], list)


def test_clean_pii_masked():
    r = client.post("/api/v1/clean", json={"text": "联系 13812345678 或 service@example.com"})
    body = r.json()
    data = body["data"]
    assert "138****5678" in data["cleaned_text"]
    assert "s***@example.com" in data["cleaned_text"]
    assert "pii_masked" in data["removed_flags"]


def test_parse_text_fallback():
    content = "七天无理由退货：签收七日内可退。\n\n质量问题：支持换货。".encode("utf-8")
    r = client.post(
        "/api/v1/parse",
        files={"file": ("policy.txt", content, "text/plain")},
        data={"options": '{"ocr": false, "clean": true}'},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True
    data = body["data"]
    assert len(data["blocks"]) >= 1
    assert all(b["block_type"] for b in data["blocks"])
    assert 0.0 <= data["clean_score"] <= 1.0
    assert "text_fallback" in data["flags"]


def test_rerank_degraded_without_model():
    r = client.post(
        "/api/v1/rerank",
        json={
            "query": "七天无理由退货",
            "documents": [
                {"id": "1", "text": "签收七日内可无理由退货"},
                {"id": "2", "text": "物流查询方式"},
            ],
            "top_n": 2,
        },
    )
    assert r.status_code == 200
    body = r.json()
    # 未装模型时：ok=false 且带 error（Java 据此降级）
    assert body["ok"] is False
    assert body["error"]


# ---------------------------------------------------------------------------
# 多格式解析（docx / xlsx / md）
# ---------------------------------------------------------------------------
def _docx_bytes() -> bytes:
    import io

    import docx
    d = docx.Document()
    d.add_heading("售后政策", level=1)
    d.add_paragraph("自签收之日起七日内可无理由退货。")
    d.add_heading("退货流程", level=2)
    d.add_paragraph("用户在订单详情页申请售后。")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "退货时效"
    table.cell(1, 1).text = "1-3 个工作日"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _xlsx_bytes() -> bytes:
    import io

    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "退款规则"
    ws.append(["规则", "说明"])
    ws.append(["退款时效", "1-3 个工作日"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_parse_docx_structured():
    docx = pytest.importorskip("docx")
    r = client.post(
        "/api/v1/parse",
        files={"file": ("policy.docx", _docx_bytes(), "application/octet-stream")},
        data={"options": '{"ocr": false, "clean": false}'},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True
    blocks = body["data"]["blocks"]
    types = {b["block_type"] for b in blocks}
    assert "heading" in types
    assert "text" in types
    assert "table" in types
    assert body["data"]["flags"] == ["docx"]
    # 标题文本带 # 前缀表层级
    headings = [b for b in blocks if b["block_type"] == "heading"]
    assert any(b["text"].startswith("# ") for b in headings)


def test_parse_xlsx_structured():
    pytest.importorskip("openpyxl")
    r = client.post(
        "/api/v1/parse",
        files={"file": ("refund.xlsx", _xlsx_bytes(), "application/octet-stream")},
        data={"options": '{"ocr": false, "clean": false}'},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True
    blocks = body["data"]["blocks"]
    types = {b["block_type"] for b in blocks}
    assert "section" in types      # 工作表名
    assert "table" in types        # 单元格数据
    assert body["data"]["flags"] == ["xlsx"]


def test_parse_markdown_keeps_syntax():
    md = (
        "# 售后政策\n\n"
        "## 七天无理由退货\n\n"
        "- 签收七日内可退\n"
        "- 商品需保持完好\n\n"
        "| 项目 | 说明 |\n| --- | --- |\n| 时效 | 1-3 天 |\n"
    ).encode("utf-8")
    r = client.post(
        "/api/v1/parse",
        files={"file": ("policy.md", md, "text/markdown")},
        data={"options": '{"ocr": false, "clean": false}'},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True
    blocks = body["data"]["blocks"]
    types = {b["block_type"] for b in blocks}
    assert "heading" in types
    assert "list" in types
    assert "table" in types
    # 保留 markdown 语法标记
    all_text = "\n".join(b["text"] for b in blocks)
    assert "# 售后政策" in all_text
    assert "- 签收七日内可退" in all_text
    assert "| 项目 | 说明 |" in all_text


def test_parse_corrupted_file_friendly_error():
    r = client.post(
        "/api/v1/parse",
        files={"file": ("broken.docx", b"not-a-real-docx-bytes", "application/octet-stream")},
        data={"options": '{"ocr": false, "clean": false}'},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is False
    assert body["error"]


# ---------------------------------------------------------------------------
# 导出（txt / md / docx / xlsx / pdf）
# ---------------------------------------------------------------------------
def test_export_txt_and_md():
    for fmt, ctype in (("txt", "text/plain"), ("md", "text/markdown")):
        r = client.post("/api/v1/export", json={"format": fmt, "text": "售后政策\n七天无理由退货"})
        assert r.status_code == 200
        assert r.headers["content-type"].startswith(ctype)
        assert "七天无理由退货" in r.content.decode("utf-8")


def test_export_docx_roundtrip():
    pytest.importorskip("docx")
    r = client.post("/api/v1/export", json={"format": "docx", "text": "售后政策\n\n七天无理由退货"})
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats")
    assert len(r.content) > 0


def test_export_xlsx_roundtrip():
    pytest.importorskip("openpyxl")
    r = client.post("/api/v1/export", json={"format": "xlsx", "text": "规则\t说明\n退款时效\t1-3天"})
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument.spreadsheetml")
    assert len(r.content) > 0


def test_export_unsupported_format():
    r = client.post("/api/v1/export", json={"format": "pptx", "text": "x"})
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is False
    assert body["error"]


def test_export_pdf():
    pytest.importorskip("fpdf")
    r = client.post("/api/v1/export", json={"format": "pdf", "text": "售后政策\n七天无理由退货"})
    assert r.status_code == 200
    if r.headers["content-type"] == "application/json":
        # 无系统中文字体时返回友好错误（避免乱码）
        body = r.json()
        assert body["ok"] is False
        assert body["error"]
    else:
        assert r.headers["content-type"] == "application/pdf"
        assert r.content.startswith(b"%PDF")
